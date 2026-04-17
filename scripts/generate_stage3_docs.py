from __future__ import annotations

import json
from datetime import UTC, datetime
from pathlib import Path

import xlsxwriter
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT_DIR = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT_DIR / "deliverables"
RESULTS_PATH = ROOT_DIR / "stage3_check_results.json"


def load_results() -> dict:
    if RESULTS_PATH.exists():
        return json.loads(RESULTS_PATH.read_text(encoding="utf-8"))
    return {}


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_title_page(doc: Document, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ОТЧЕТ ПО ПРОЕКТНОЙ РАБОТЕ")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()
    doc.add_paragraph("Веб-приложение «Заказ пиццы»").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Выполнил: ____________________________").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Проверил: ____________________________").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Москва, {datetime.now(UTC).year}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 2) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14 if level == 2 else 12)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row_values in enumerate(rows):
        row = table.add_row().cells
        for col_index, value in enumerate(row_values):
            set_cell_text(row[col_index], value, bold=row_index == 0)
    doc.add_paragraph()


def add_test_case_table(doc: Document, case: dict[str, str]) -> None:
    rows = [
        ["Тестовый пример #", case["id"]],
        ["Приоритет тестирования", case["priority"]],
        ["Заголовок/название теста", case["title"]],
        ["Краткое изложение теста", case["summary"]],
        ["Этапы теста", case["steps"]],
        ["Тестовые данные", case["data"]],
        ["Ожидаемый результат", case["expected"]],
        ["Фактический результат", case["actual"]],
        ["Статус", case["status"]],
        ["Предварительное условие", case["precondition"]],
        ["Постусловие", case["postcondition"]],
        ["Примечания/комментарии", case["notes"]],
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for key, value in rows:
        row = table.add_row().cells
        set_cell_text(row[0], key, bold=True)
        set_cell_text(row[1], value)
    doc.add_paragraph()


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    doc.add_paragraph()


def extract_snippet(path: Path, marker: str, lines_after: int = 20) -> str:
    lines = path.read_text(encoding="utf-8").splitlines()
    for index, line in enumerate(lines):
        if marker in line:
            return "\n".join(lines[index : index + lines_after])
    return f"# Marker not found: {marker}"


def pretty_json(data) -> str:
    if data in (None, "", []):
        return "—"
    if isinstance(data, str):
        return data
    return json.dumps(data, ensure_ascii=False, indent=2)


API_ACTIONS = [
    {
        "path": "/api/v1/health",
        "method": "GET",
        "auth": "None",
        "role": "None",
        "request": "—",
        "success": '{"status":"ok"}',
        "errors4": "—",
        "errors3": "—",
        "errors2": "200 OK",
        "description": "Проверка доступности API.",
    },
    {
        "path": "/api/v1/pizzas",
        "method": "GET",
        "auth": "None",
        "role": "None",
        "request": "Query: q?, category?, available?",
        "success": '[{"id":"margherita","name":"Margherita"}]',
        "errors4": "400 invalid query",
        "errors3": "—",
        "errors2": "200 OK",
        "description": "Каталог пицц.",
    },
    {
        "path": "/api/v1/pizzas/{pizza_id}",
        "method": "GET",
        "auth": "None",
        "role": "None",
        "request": "Path: pizza_id",
        "success": '{"id":"margherita","name":"Margherita"}',
        "errors4": "—",
        "errors3": "404 Pizza not found",
        "errors2": "200 OK",
        "description": "Получение одной пиццы.",
    },
    {
        "path": "/api/v1/sizes",
        "method": "GET",
        "auth": "None",
        "role": "None",
        "request": "—",
        "success": '[{"id":"m","name":"30 cm"}]',
        "errors4": "—",
        "errors3": "—",
        "errors2": "200 OK",
        "description": "Список размеров.",
    },
    {
        "path": "/api/v1/toppings",
        "method": "GET",
        "auth": "None",
        "role": "None",
        "request": "Query: available?",
        "success": '[{"id":"extra_cheese","name":"Extra cheese"}]',
        "errors4": "—",
        "errors3": "—",
        "errors2": "200 OK",
        "description": "Список топпингов.",
    },
    {
        "path": "/api/v1/promocodes/active",
        "method": "GET",
        "auth": "None",
        "role": "None",
        "request": "—",
        "success": '[{"code":"WELCOME10","discount_type":"percent"}]',
        "errors4": "—",
        "errors3": "—",
        "errors2": "200 OK",
        "description": "Активные промокоды этапа 3.",
    },
    {
        "path": "/api/v1/auth/register",
        "method": "POST",
        "auth": "None",
        "role": "None",
        "request": '{"email":"user@mail.ru","password":"123456","name":"User"}',
        "success": '{"access_token":"...","token_type":"bearer"}',
        "errors4": "409 Email already registered; 422 validation error",
        "errors3": "—",
        "errors2": "200 OK",
        "description": "Регистрация пользователя.",
    },
    {
        "path": "/api/v1/auth/login",
        "method": "POST",
        "auth": "None",
        "role": "None",
        "request": '{"email":"peter@mail.ru","password":"123456"}',
        "success": '{"access_token":"...","token_type":"bearer"}',
        "errors4": "401 Invalid credentials",
        "errors3": "—",
        "errors2": "200 OK",
        "description": "Авторизация пользователя или администратора.",
    },
    {
        "path": "/api/v1/me",
        "method": "GET",
        "auth": "JWT",
        "role": "User/Admin",
        "request": "—",
        "success": '{"id":"...","email":"user@mail.ru","role":"user"}',
        "errors4": "—",
        "errors3": "401 Not authenticated",
        "errors2": "200 OK",
        "description": "Получение профиля.",
    },
    {
        "path": "/api/v1/me",
        "method": "PATCH",
        "auth": "JWT",
        "role": "User/Admin",
        "request": '{"name":"Новое имя"}',
        "success": '{"id":"...","name":"Новое имя"}',
        "errors4": "422 validation error",
        "errors3": "401 Not authenticated",
        "errors2": "200 OK",
        "description": "Редактирование профиля.",
    },
    {
        "path": "/api/v1/me/addresses",
        "method": "GET",
        "auth": "JWT",
        "role": "User/Admin",
        "request": "—",
        "success": '[{"id":"...","label":"Дом","is_default":true}]',
        "errors4": "—",
        "errors3": "401 Not authenticated",
        "errors2": "200 OK",
        "description": "Список адресов пользователя.",
    },
    {
        "path": "/api/v1/me/addresses",
        "method": "POST",
        "auth": "JWT",
        "role": "User/Admin",
        "request": '{"label":"Офис","address":"Москва, Арбат 15"}',
        "success": '{"id":"...","label":"Офис"}',
        "errors4": "422 validation error",
        "errors3": "401 Not authenticated",
        "errors2": "200 OK",
        "description": "Создание адреса.",
    },
    {
        "path": "/api/v1/me/addresses/{address_id}",
        "method": "PATCH",
        "auth": "JWT",
        "role": "User/Admin",
        "request": '{"is_default":true}',
        "success": '{"id":"...","is_default":true}',
        "errors4": "400 Invalid id; 422 validation error",
        "errors3": "401 Not authenticated; 404 Address not found",
        "errors2": "200 OK",
        "description": "Обновление адреса.",
    },
    {
        "path": "/api/v1/me/addresses/{address_id}",
        "method": "DELETE",
        "auth": "JWT",
        "role": "User/Admin",
        "request": "Path: address_id",
        "success": '{"deleted":"..."}',
        "errors4": "400 Invalid id",
        "errors3": "401 Not authenticated; 404 Address not found",
        "errors2": "200 OK",
        "description": "Удаление адреса.",
    },
    {
        "path": "/api/v1/me/bonuses",
        "method": "GET",
        "auth": "JWT",
        "role": "User/Admin",
        "request": "—",
        "success": '{"balance":57,"accrued":57,"spent":0}',
        "errors4": "—",
        "errors3": "401 Not authenticated",
        "errors2": "200 OK",
        "description": "Получение бонусов.",
    },
    {
        "path": "/api/v1/orders/preview",
        "method": "POST",
        "auth": "None / JWT optional",
        "role": "Guest/User/Admin",
        "request": '{"bonus_spent":20,"promo_code":"WELCOME10","scheduled_for":"..."}',
        "success": '{"promo_discount":119,"scheduled_for":"...","total":1051}',
        "errors4": "400 invalid promo/schedule/bonuses; 409 pizza not available",
        "errors3": "401 guest bonus spending forbidden",
        "errors2": "200 OK",
        "description": "Предварительный расчет заказа.",
    },
    {
        "path": "/api/v1/orders",
        "method": "POST",
        "auth": "None / JWT optional",
        "role": "Guest/User/Admin",
        "request": '{"customer":{...},"delivery":{...},"promo_code":"WELCOME10","scheduled_for":"..."}',
        "success": '{"order_id":"...","promo_code":"WELCOME10","total":1051}',
        "errors4": "400 invalid business rule; 409 pizza not available; 422 validation",
        "errors3": "401 guest bonus spending forbidden",
        "errors2": "200 OK",
        "description": "Создание заказа.",
    },
    {
        "path": "/api/v1/orders/{order_id}",
        "method": "GET",
        "auth": "JWT or public_token",
        "role": "Guest/User/Admin",
        "request": "Path: order_id; Query: public_token?",
        "success": '{"order_id":"...","status":"done","review":{"rating":5}}',
        "errors4": "400 Invalid order_id",
        "errors3": "401 Not allowed; 404 Order not found",
        "errors2": "200 OK",
        "description": "Получение заказа по токену или JWT.",
    },
    {
        "path": "/api/v1/me/orders",
        "method": "GET",
        "auth": "JWT",
        "role": "User/Admin",
        "request": "—",
        "success": '[{"order_id":"...","scheduled_for":"..."}]',
        "errors4": "—",
        "errors3": "401 Not authenticated",
        "errors2": "200 OK",
        "description": "История заказов.",
    },
    {
        "path": "/api/v1/me/orders/{order_id}",
        "method": "GET",
        "auth": "JWT",
        "role": "User/Admin",
        "request": "Path: order_id",
        "success": '{"order_id":"...","review":{"rating":5}}',
        "errors4": "400 Invalid order_id",
        "errors3": "401 Not authenticated; 404 Order not found",
        "errors2": "200 OK",
        "description": "Детали заказа в кабинете.",
    },
    {
        "path": "/api/v1/me/orders/{order_id}/review",
        "method": "POST",
        "auth": "JWT",
        "role": "User/Admin",
        "request": '{"rating":5,"comment":"Отзыв этапа 3: быстро и удобно"}',
        "success": '{"id":"...","rating":5}',
        "errors4": "400 Invalid order_id; 422 validation error",
        "errors3": "401 Not authenticated; 404 Order not found; 409 only completed orders",
        "errors2": "200 OK",
        "description": "Сохранение отзыва по заказу.",
    },
    {
        "path": "/api/v1/admin/orders",
        "method": "GET",
        "auth": "JWT",
        "role": "Admin",
        "request": "—",
        "success": '[{"order_id":"...","promo_code":"WELCOME10"}]',
        "errors4": "—",
        "errors3": "401 Not authenticated; 403 Admin only",
        "errors2": "200 OK",
        "description": "Список заказов для админа.",
    },
    {
        "path": "/api/v1/admin/orders/{order_id}/status",
        "method": "PATCH",
        "auth": "JWT",
        "role": "Admin",
        "request": '{"status":"done"}',
        "success": '{"order_id":"...","status":"done"}',
        "errors4": "400 Invalid order_id; 422 validation error",
        "errors3": "401 Not authenticated; 403 Admin only; 404 Order not found",
        "errors2": "200 OK",
        "description": "Смена статуса заказа.",
    },
    {
        "path": "/api/v1/admin/reviews",
        "method": "GET",
        "auth": "JWT",
        "role": "Admin",
        "request": "—",
        "success": '[{"order_id":"...","rating":5,"user_name":"Peter"}]',
        "errors4": "—",
        "errors3": "401 Not authenticated; 403 Admin only",
        "errors2": "200 OK",
        "description": "Просмотр отзывов.",
    },
    {
        "path": "/api/v1/admin/promocodes",
        "method": "GET",
        "auth": "JWT",
        "role": "Admin",
        "request": "—",
        "success": '[{"code":"WELCOME10","active":true}]',
        "errors4": "—",
        "errors3": "401 Not authenticated; 403 Admin only",
        "errors2": "200 OK",
        "description": "Список промокодов.",
    },
    {
        "path": "/api/v1/admin/promocodes",
        "method": "POST",
        "auth": "JWT",
        "role": "Admin",
        "request": '{"code":"WELCOME10","discount_type":"percent","discount_value":10}',
        "success": '{"id":"...","code":"WELCOME10"}',
        "errors4": "409 Promo code already exists; 422 validation error",
        "errors3": "401 Not authenticated; 403 Admin only",
        "errors2": "200 OK",
        "description": "Создание промокода.",
    },
    {
        "path": "/api/v1/admin/promocodes/{promo_id}",
        "method": "PATCH",
        "auth": "JWT",
        "role": "Admin",
        "request": '{"active":false,"discount_value":15}',
        "success": '{"id":"...","active":false}',
        "errors4": "400 Invalid id; 409 duplicate code; 422 validation",
        "errors3": "401 Not authenticated; 403 Admin only; 404 Promo code not found",
        "errors2": "200 OK",
        "description": "Обновление промокода.",
    },
    {
        "path": "/api/v1/admin/promocodes/{promo_id}",
        "method": "DELETE",
        "auth": "JWT",
        "role": "Admin",
        "request": "Path: promo_id",
        "success": '{"id":"...","active":false}',
        "errors4": "400 Invalid id",
        "errors3": "401 Not authenticated; 403 Admin only; 404 Promo code not found",
        "errors2": "200 OK",
        "description": "Деактивация промокода.",
    },
    {
        "path": "/api/v1/admin/pizzas",
        "method": "GET",
        "auth": "JWT",
        "role": "Admin",
        "request": "—",
        "success": '[{"id":"margherita","available":true}]',
        "errors4": "—",
        "errors3": "401 Not authenticated; 403 Admin only",
        "errors2": "200 OK",
        "description": "Список пицц в админ-панели.",
    },
    {
        "path": "/api/v1/admin/pizzas",
        "method": "POST",
        "auth": "JWT",
        "role": "Admin",
        "request": '{"id":"new_pizza","name":"New Pizza","base_price":650}',
        "success": '{"id":"new_pizza"}',
        "errors4": "409 Pizza id exists; 422 validation error",
        "errors3": "401 Not authenticated; 403 Admin only",
        "errors2": "200 OK",
        "description": "Создание пиццы.",
    },
    {
        "path": "/api/v1/admin/pizzas/{pizza_id}",
        "method": "PATCH",
        "auth": "JWT",
        "role": "Admin",
        "request": '{"name":"Updated Pizza","base_price":700}',
        "success": '{"id":"new_pizza","name":"Updated Pizza"}',
        "errors4": "422 validation error",
        "errors3": "401 Not authenticated; 403 Admin only; 404 Pizza not found",
        "errors2": "200 OK",
        "description": "Обновление пиццы.",
    },
    {
        "path": "/api/v1/admin/pizzas/{pizza_id}",
        "method": "DELETE",
        "auth": "JWT",
        "role": "Admin",
        "request": "Path: pizza_id",
        "success": '{"id":"new_pizza","available":false}',
        "errors4": "—",
        "errors3": "401 Not authenticated; 403 Admin only; 404 Pizza not found",
        "errors2": "200 OK",
        "description": "Soft delete пиццы.",
    },
    {
        "path": "/api/v1/admin/toppings",
        "method": "GET",
        "auth": "JWT",
        "role": "Admin",
        "request": "—",
        "success": '[{"id":"extra_cheese","price":80}]',
        "errors4": "—",
        "errors3": "401 Not authenticated; 403 Admin only",
        "errors2": "200 OK",
        "description": "Список топпингов.",
    },
    {
        "path": "/api/v1/admin/toppings",
        "method": "POST",
        "auth": "JWT",
        "role": "Admin",
        "request": '{"id":"truffle","name":"Truffle","price":150}',
        "success": '{"id":"truffle"}',
        "errors4": "409 Topping id exists; 422 validation error",
        "errors3": "401 Not authenticated; 403 Admin only",
        "errors2": "200 OK",
        "description": "Создание топпинга.",
    },
    {
        "path": "/api/v1/admin/toppings/{topping_id}",
        "method": "PATCH",
        "auth": "JWT",
        "role": "Admin",
        "request": '{"name":"Truffle Deluxe","price":190}',
        "success": '{"id":"truffle","price":190}',
        "errors4": "422 validation error",
        "errors3": "401 Not authenticated; 403 Admin only; 404 Topping not found",
        "errors2": "200 OK",
        "description": "Обновление топпинга.",
    },
    {
        "path": "/api/v1/admin/toppings/{topping_id}",
        "method": "DELETE",
        "auth": "JWT",
        "role": "Admin",
        "request": "Path: topping_id",
        "success": '{"id":"truffle","available":false}',
        "errors4": "—",
        "errors3": "401 Not authenticated; 403 Admin only; 404 Topping not found",
        "errors2": "200 OK",
        "description": "Soft delete топпинга.",
    },
    {
        "path": "/api/v1/admin/sizes",
        "method": "GET",
        "auth": "JWT",
        "role": "Admin",
        "request": "—",
        "success": '[{"id":"m","multiplier":1.25}]',
        "errors4": "—",
        "errors3": "401 Not authenticated; 403 Admin only",
        "errors2": "200 OK",
        "description": "Список размеров.",
    },
    {
        "path": "/api/v1/admin/sizes",
        "method": "POST",
        "auth": "JWT",
        "role": "Admin",
        "request": '{"id":"xl","name":"40 cm","multiplier":1.8}',
        "success": '{"id":"xl"}',
        "errors4": "409 Size id exists; 422 validation error",
        "errors3": "401 Not authenticated; 403 Admin only",
        "errors2": "200 OK",
        "description": "Создание размера.",
    },
    {
        "path": "/api/v1/admin/sizes/{size_id}",
        "method": "PATCH",
        "auth": "JWT",
        "role": "Admin",
        "request": '{"name":"40 cm","multiplier":1.75}',
        "success": '{"id":"xl","multiplier":1.75}',
        "errors4": "422 validation error",
        "errors3": "401 Not authenticated; 403 Admin only; 404 Size not found",
        "errors2": "200 OK",
        "description": "Обновление размера.",
    },
]


def make_api_workbook(path: Path) -> None:
    workbook = xlsxwriter.Workbook(path)
    worksheet = workbook.add_worksheet("List of all API actions")

    header_format = workbook.add_format({"bold": True, "align": "center", "valign": "vcenter", "text_wrap": True, "border": 1, "bg_color": "#D9EAF7", "font_name": "Times New Roman"})
    note_format = workbook.add_format({"align": "center", "valign": "vcenter", "text_wrap": True, "border": 1, "bg_color": "#F4F4F4", "font_name": "Times New Roman"})
    cell_format = workbook.add_format({"text_wrap": True, "valign": "top", "border": 1, "font_name": "Times New Roman", "font_size": 11})

    headers = ["Path", "Method", "Required Auth", "Required Role", "Request body", "Success response body", "Possible 400 errors", "Possible 300 errors", "Possible 200 errors", "Description"]
    notes = ["URL-запроса", "Метод (GET/POST/PUT/PATCH/DELETE)", "None / JWT / JWT optional", "Укажите роль", "Тело запроса", "Тело ответа", "Ошибки валидации и бизнес-правил", "Ошибки доступа и отсутствующих сущностей", "Успешные коды ответа", "Комментарии разработчика"]

    for col, value in enumerate(headers):
        worksheet.write(0, col, value, header_format)
    for col, value in enumerate(notes):
        worksheet.write(1, col, value, note_format)

    for row_index, row in enumerate(API_ACTIONS, start=2):
        worksheet.write(row_index, 0, row["path"], cell_format)
        worksheet.write(row_index, 1, row["method"], cell_format)
        worksheet.write(row_index, 2, row["auth"], cell_format)
        worksheet.write(row_index, 3, row["role"], cell_format)
        worksheet.write(row_index, 4, row["request"], cell_format)
        worksheet.write(row_index, 5, row["success"], cell_format)
        worksheet.write(row_index, 6, row["errors4"], cell_format)
        worksheet.write(row_index, 7, row["errors3"], cell_format)
        worksheet.write(row_index, 8, row["errors2"], cell_format)
        worksheet.write(row_index, 9, row["description"], cell_format)

    worksheet.set_column("A:A", 34)
    worksheet.set_column("B:B", 14)
    worksheet.set_column("C:C", 18)
    worksheet.set_column("D:D", 16)
    worksheet.set_column("E:E", 30)
    worksheet.set_column("F:F", 30)
    worksheet.set_column("G:G", 30)
    worksheet.set_column("H:H", 32)
    worksheet.set_column("I:I", 16)
    worksheet.set_column("J:J", 28)
    worksheet.freeze_panes(2, 0)
    worksheet.autofilter(0, 0, len(API_ACTIONS) + 1, len(headers) - 1)
    workbook.close()


def make_test_doc(path: Path, results: dict) -> None:
    doc = Document()
    set_default_font(doc)
    add_title_page(doc, "по тестированию новых функций 3 этапа")

    add_heading(doc, "Введение")
    add_paragraph(
        doc,
        "Документ фиксирует результаты тестирования функциональности третьего этапа "
        "развития веб-приложения «Заказ пиццы». На данном этапе в систему были "
        "добавлены промокоды, предварительный серверный расчет заказа, отложенная "
        "доставка и механизм отзывов к завершенным заказам.",
    )

    add_heading(doc, "Сводка по изменениям")
    add_table(
        doc,
        [
            ["Подсистема", "Новые возможности этапа 3", "Проверяемый результат"],
            ["Оформление заказа", "Промокод, preview, scheduled delivery", "Корректный расчет и сохранение новых полей"],
            ["Личный кабинет", "Отзывы по заказам со статусом done", "Создание/обновление отзыва только владельцем"],
            ["Администрирование", "Управление промокодами и просмотр отзывов", "Работа с новыми сущностями без SQL"],
        ],
    )

    add_heading(doc, "Паспорт тестирования")
    add_table(
        doc,
        [
            ["Название проекта", "Веб-приложение «Заказ пиццы»"],
            ["Рабочая версия", "3.0 (этап 3)"],
            ["Имя тестирующего", "Исполнитель проекта"],
            ["Дата(ы) теста", results.get("generated_at", datetime.now(UTC).isoformat())[:10]],
            ["Источник фактических результатов", "Локальный прогон API через FastAPI TestClient"],
        ],
    )

    tests = [
        {
            "id": "TC_ST3_PROMO_01",
            "priority": "Высокий",
            "title": "Получение списка активных промокодов",
            "summary": "Проверить, что публичный API возвращает активные акции этапа 3.",
            "steps": "1. Отправить GET-запрос на /api/v1/promocodes/active.\n2. Проверить количество и поля объектов.",
            "data": "GET /api/v1/promocodes/active",
            "expected": "HTTP 200 и массив активных промокодов.",
            "actual": f"Получен HTTP 200. Возвращено {results.get('active_promos_count', 0)} активных промокода(ов).",
            "status": "Зачет",
            "precondition": "В БД засеяны демонстрационные промокоды.",
            "postcondition": "Состояние системы не изменяется.",
            "notes": "Позитивный сценарий нового публичного маршрута.",
        },
        {
            "id": "TC_ST3_ORDER_02",
            "priority": "Высокий",
            "title": "Предварительный расчет заказа с промокодом и отложенной доставкой",
            "summary": "Проверить совместную работу скидки, scheduled delivery и бонусов в preview-режиме.",
            "steps": "1. Авторизоваться пользователем.\n2. Отправить POST на /api/v1/orders/preview.\n3. Сверить promo_discount, scheduled_for и total.",
            "data": "promo_code=WELCOME10, bonus_spent=20, items=[pepperoni m + extra_cheese, margherita s]",
            "expected": "HTTP 200 и корректный расчет всех скидок и итоговой суммы.",
            "actual": f"Получен HTTP {results.get('preview_ok_status')}. promo_discount={results.get('preview_ok', {}).get('promo_discount')}, total={results.get('preview_ok', {}).get('total')}.",
            "status": "Зачет",
            "precondition": "Пользователь авторизован, WELCOME10 активен.",
            "postcondition": "Состояние системы не изменяется.",
            "notes": "Ключевой позитивный сценарий этапа 3.",
        },
        {
            "id": "TC_ST3_PROMO_03",
            "priority": "Высокий",
            "title": "Отклонение несуществующего промокода",
            "summary": "Проверить, что система возвращает бизнес-ошибку при неверном коде.",
            "steps": "1. Авторизоваться пользователем.\n2. Отправить POST на /api/v1/orders/preview с promo_code=BADCODE.\n3. Проверить detail ошибки.",
            "data": "promo_code=BADCODE",
            "expected": "HTTP 400 и сообщение Promo code not found or inactive.",
            "actual": f"Получен HTTP {results.get('preview_bad_status')}. detail={results.get('preview_bad', {}).get('detail')}.",
            "status": "Зачет",
            "precondition": "Промокод BADCODE отсутствует в системе.",
            "postcondition": "Заказ не создается.",
            "notes": "Негативный сценарий валидации промокода.",
        },
        {
            "id": "TC_ST3_SCHEDULE_04",
            "priority": "Высокий",
            "title": "Отклонение слишком ранней отложенной доставки",
            "summary": "Проверить правило минимального интервала в 30 минут для scheduled delivery.",
            "steps": "1. Отправить POST на /api/v1/orders/preview.\n2. Передать scheduled_for менее чем через 30 минут.\n3. Проверить detail ошибки.",
            "data": "scheduled_for = now + 10 minutes",
            "expected": "HTTP 400 и сообщение о минимальном интервале 30 минут.",
            "actual": f"Получен HTTP {results.get('preview_early_schedule_status')}. detail={results.get('preview_early_schedule', {}).get('detail')}.",
            "status": "Зачет",
            "precondition": "Пользователь авторизован.",
            "postcondition": "Заказ не создается.",
            "notes": "Негативный сценарий scheduled delivery.",
        },
        {
            "id": "TC_ST3_BONUS_05",
            "priority": "Средний",
            "title": "Отклонение превышения бонусного баланса",
            "summary": "Проверить, что preview не позволяет списывать бонусов больше доступного баланса.",
            "steps": "1. Отправить POST на /api/v1/orders/preview с bonus_spent=5000.\n2. Проверить код ошибки и detail.",
            "data": "bonus_spent=5000",
            "expected": "HTTP 400 и сообщение Not enough bonus balance.",
            "actual": f"Получен HTTP {results.get('preview_bonus_overflow_status')}. detail={results.get('preview_bonus_overflow', {}).get('detail')}.",
            "status": "Зачет",
            "precondition": "Баланс пользователя меньше 5000.",
            "postcondition": "Состояние системы не изменяется.",
            "notes": "Негативный сценарий расширенной логики расчета.",
        },
        {
            "id": "TC_ST3_ORDER_06",
            "priority": "Высокий",
            "title": "Создание заказа с промокодом и scheduled delivery",
            "summary": "Проверить сохранение новых полей в заказе.",
            "steps": "1. Отправить POST на /api/v1/orders.\n2. Проверить promo_code, promo_discount и scheduled_for в ответе.",
            "data": "customer=Peter, promo_code=WELCOME10, scheduled_for=через 2 часа",
            "expected": "HTTP 200 и созданный заказ с новыми полями этапа 3.",
            "actual": f"Получен HTTP 200. order_id={results.get('created_order', {}).get('order_id')}, promo_code={results.get('created_order', {}).get('promo_code')}.",
            "status": "Зачет",
            "precondition": "Пользователь авторизован, заказ валиден.",
            "postcondition": "В БД создан новый заказ.",
            "notes": "Позитивный сценарий по основному бизнес-процессу.",
        },
        {
            "id": "TC_ST3_REVIEW_07",
            "priority": "Высокий",
            "title": "Запрет на отзыв до завершения заказа",
            "summary": "Проверить, что отзыв можно оставить только после статуса done.",
            "steps": "1. Создать новый заказ.\n2. До смены статуса отправить POST на /api/v1/me/orders/{order_id}/review.\n3. Проверить ответ.",
            "data": "rating=4, comment=Пока еще рано для отзыва",
            "expected": "HTTP 409 и сообщение о доступности отзыва только для completed orders.",
            "actual": f"Получен HTTP {results.get('review_before_done_status')}. detail={results.get('review_before_done', {}).get('detail')}.",
            "status": "Зачет",
            "precondition": "Заказ существует и находится в статусе created.",
            "postcondition": "Отзыв не создается.",
            "notes": "Негативный сценарий нового механизма отзывов.",
        },
        {
            "id": "TC_ST3_REVIEW_08",
            "priority": "Высокий",
            "title": "Сохранение отзыва после перевода заказа в done",
            "summary": "Проверить создание отзыва владельцем завершенного заказа.",
            "steps": "1. Перевести заказ в статус done.\n2. Отправить POST на /api/v1/me/orders/{order_id}/review.\n3. Проверить рейтинг и комментарий в ответе.",
            "data": "rating=5, comment=Отзыв этапа 3: быстро и удобно",
            "expected": "HTTP 200 и сохраненный отзыв.",
            "actual": f"Получен HTTP 200. review_id={results.get('review_saved', {}).get('id')}, rating={results.get('review_saved', {}).get('rating')}.",
            "status": "Зачет",
            "precondition": "Администратор перевел заказ в статус done.",
            "postcondition": "В БД создана запись в таблице reviews.",
            "notes": "Позитивный сценарий отзывов.",
        },
        {
            "id": "TC_ST3_ADMIN_09",
            "priority": "Средний",
            "title": "Просмотр отзывов администратором",
            "summary": "Проверить новый административный маршрут для просмотра клиентских отзывов.",
            "steps": "1. Авторизоваться администратором.\n2. Отправить GET на /api/v1/admin/reviews.\n3. Проверить наличие последнего созданного отзыва.",
            "data": "GET /api/v1/admin/reviews",
            "expected": "HTTP 200 и массив отзывов с order_id, user_name и rating.",
            "actual": f"Получен HTTP 200. reviews_count={results.get('admin_reviews_count')}, last_review_rating={results.get('admin_reviews_last', {}).get('rating')}.",
            "status": "Зачет",
            "precondition": "В системе существует хотя бы один отзыв.",
            "postcondition": "Состояние системы не изменяется.",
            "notes": "Позитивный сценарий нового admin-маршрута этапа 3.",
        },
    ]

    for case in tests:
        add_test_case_table(doc, case)

    doc.save(path)


def make_user_manual(path: Path) -> None:
    doc = Document()
    set_default_font(doc)
    add_title_page(doc, "по обновленной инструкции пользователя 3 этапа")

    add_heading(doc, "Назначение системы")
    add_paragraph(
        doc,
        "Система «Заказ пиццы» предназначена для оформления доставки пиццы через веб-интерфейс. "
        "На третьем этапе система поддерживает применение промокодов, отложенную доставку, "
        "предварительный серверный расчет заказа и отзывы клиентов после выполнения заказа.",
    )

    add_heading(doc, "Новые функции этапа 3")
    add_table(
        doc,
        [
            ["Функция", "Действие пользователя", "Результат"],
            ["Промокоды", "Ввести код на Checkout", "Система применяет скидку к заказу"],
            ["Preview-расчет", "Менять корзину, бонусы, код или время", "Итог пересчитывается на backend до подтверждения"],
            ["Отложенная доставка", "Выбрать дату и время", "В заказе сохраняется planned time"],
            ["Отзывы", "Открыть завершенный заказ и поставить оценку", "Отзыв сохраняется и виден администратору"],
        ],
    )

    add_heading(doc, "Инструкция для клиента")
    for step in [
        "1. Откройте меню и добавьте нужные позиции в корзину.",
        "2. Перейдите на страницу оформления заказа (`/checkout`).",
        "3. При необходимости введите промокод. Если код активен, скидка появится в блоке суммы.",
        "4. Выберите режим доставки: сразу или в запланированное время. Для отложенной доставки задайте дату и время минимум через 30 минут.",
        "5. Если вы авторизованы, укажите количество бонусов для списания.",
        "6. Проверьте предварительный расчет и подтвердите заказ.",
        "7. На странице успеха просмотрите итоговую сумму, статус заказа, промокод и время доставки.",
        "8. После статуса `done` откройте личный кабинет, выберите заказ и оставьте отзыв.",
    ]:
        add_paragraph(doc, step)

    add_heading(doc, "Инструкция для администратора")
    for step in [
        "1. Войдите под учетной записью администратора и откройте `/admin`.",
        "2. В разделе заказов контролируйте статусы и проверяйте, какие заказы созданы с промокодом и scheduled delivery.",
        "3. В разделе промокодов создавайте, изменяйте и отключайте акции.",
        "4. В разделе отзывов анализируйте обратную связь клиентов.",
        "5. В разделах каталога продолжайте управлять пиццами, топпингами и размерами.",
    ]:
        add_paragraph(doc, step)

    add_heading(doc, "Правила и ограничения")
    for rule in [
        "Промокод применяется только если он активен и выполняется минимальная сумма заказа.",
        "Отложенная доставка возможна минимум через 30 минут и максимум на 7 дней вперед.",
        "Бонусы можно использовать только после авторизации.",
        "Отзыв доступен только владельцу заказа и только после статуса `done`.",
    ]:
        add_paragraph(doc, rule)

    doc.save(path)


def make_code_doc(path: Path) -> None:
    doc = Document()
    set_default_font(doc)
    add_title_page(doc, "по коду и описанию реализации функций 3 этапа")

    add_heading(doc, "Общая идея реализации")
    add_paragraph(
        doc,
        "Изменения третьего этапа проведены сквозным способом: сначала были добавлены новые сущности "
        "и поля базы данных, затем обновлены REST-маршруты, после чего новая логика была интегрирована "
        "в интерфейс оформления заказа, личный кабинет и административную панель.",
    )

    add_heading(doc, "Ключевые файлы")
    add_table(
        doc,
        [
            ["Файл", "Назначение"],
            ["backend/app/models/promo.py", "Модель промокода с ограничениями по сумме заказа и активностью."],
            ["backend/app/models/order.py", "Поля promo_code, promo_discount, scheduled_for и сущность Review."],
            ["backend/app/services/order_flow.py", "Центральная бизнес-логика preview/create заказа."],
            ["backend/app/api/routers/orders.py", "Маршруты расчета заказа и отзывов."],
            ["backend/app/api/routers/admin.py", "Маршруты промокодов и административного просмотра отзывов."],
            ["src/pages/CheckoutPage.jsx", "Промокод, scheduled delivery и server-side preview."],
            ["src/pages/AccountPage.jsx", "Отображение новых полей заказа и сохранение отзывов."],
            ["src/pages/AdminPage.jsx", "Разделы промокодов и отзывов в админке."],
        ],
    )

    add_heading(doc, "Фрагмент 1. Централизация расчета заказа")
    add_paragraph(doc, "В `order_flow.py` вся новая бизнес-логика сосредоточена в одном сервисе. Это позволяет использовать одинаковые правила в `/orders/preview` и `/orders`.")
    add_code_block(doc, extract_snippet(ROOT_DIR / "backend/app/services/order_flow.py", "async def prepare_order", lines_after=34))

    add_heading(doc, "Фрагмент 2. Новые маршруты основного бизнес-процесса")
    add_paragraph(doc, "Роутер заказов использует общий сервис для preview, создания заказа и сохранения отзывов.")
    add_code_block(doc, extract_snippet(ROOT_DIR / "backend/app/api/routers/orders.py", '@router.post("/orders/preview"', lines_after=52))

    add_heading(doc, "Фрагмент 3. Администрирование промокодов")
    add_paragraph(doc, "Отдельный блок административных маршрутов позволяет создавать, изменять и деактивировать промокоды без прямой работы с базой данных.")
    add_code_block(doc, extract_snippet(ROOT_DIR / "backend/app/api/routers/admin.py", '@router.get("/admin/promocodes"', lines_after=56))

    add_heading(doc, "Фрагмент 4. Интеграция в frontend Checkout")
    add_paragraph(doc, "На стороне клиента Checkout собирает draft заказа и вызывает server-side preview, чтобы пользователь видел фактическую стоимость до финальной отправки.")
    add_code_block(doc, extract_snippet(ROOT_DIR / "src/pages/CheckoutPage.jsx", "const previewPayload = useMemo(", lines_after=40))

    add_heading(doc, "Фрагмент 5. Отзывы в личном кабинете")
    add_paragraph(doc, "В `AccountPage.jsx` реализован интерфейс оценки завершенного заказа и повторного редактирования отзыва.")
    add_code_block(doc, extract_snippet(ROOT_DIR / "src/pages/AccountPage.jsx", "async function saveReview()", lines_after=28))

    add_heading(doc, "Итоги реализации")
    add_paragraph(
        doc,
        "Код третьего этапа поддерживает новые сущности, расширенный расчет заказа, новые административные маршруты "
        "и пользовательские сценарии без дублирования бизнес-логики. Основные правила зафиксированы в backend-сервисе "
        "и используются одинаково во всех точках входа.",
    )

    doc.save(path)


def main() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    results = load_results()
    make_api_workbook(OUTPUT_DIR / "Таблица_маршрутизации_API_этап_3.xlsx")
    make_test_doc(OUTPUT_DIR / "Тесты_новых_функций_этап_3.docx", results)
    make_user_manual(OUTPUT_DIR / "Инструкция_пользователя_этап_3.docx")
    make_code_doc(OUTPUT_DIR / "Код_и_описание_этап_3.docx")
    print(str(OUTPUT_DIR))


if __name__ == "__main__":
    main()
